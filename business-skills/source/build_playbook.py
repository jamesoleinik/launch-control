"""Build the sample launch-playbook.docx that gets extracted into Business Skills.

This is the deliberately-messy "Word doc on SharePoint nobody reads" input for
Episode 2. A coding agent reads it and emits structured Business Skill markdown.

Run from the repo root:
    python business-skills/source/build_playbook.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt


def build() -> Path:
    out = Path(__file__).parent / "launch-playbook.docx"

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Contoso Product Launch Playbook", level=0)
    title.alignment = 1

    p = doc.add_paragraph()
    p.add_run("Version 4.7  ·  Owner: PMO  ·  Last meaningful edit: ").italic = True
    p.add_run("we honestly don't remember, sometime before reorg #3").italic = True

    doc.add_paragraph(
        "This document is the official Contoso launch playbook. It is the source "
        "of truth for how we run product launches end to end. In practice nobody "
        "reads it cover to cover — it lives on SharePoint and gets opened twice a "
        "year, usually in a panic. The intent of this rewrite is to capture the "
        "policy in plain English so that a coding agent can extract it into "
        "machine-executable Business Skills. Where you see a section flagged as "
        "POLICY, that is a hard rule the agent must obey."
    )

    doc.add_heading("1. Launch Readiness — go / no-go", level=1)
    doc.add_paragraph(
        "Every launch passes through a readiness review before we ship. The PMO "
        "evaluates the launch against a set of milestone gates and produces a "
        "verdict of GO, CONDITIONAL, or NO-GO. The verdict is not a vibe — it is "
        "computed."
    )

    doc.add_heading("1.1 How readiness is computed (POLICY)", level=2)
    doc.add_paragraph(
        "Readiness is the weighted average of all milestone completion states on "
        "the launch. The PMO does NOT hand-tally gates in a slide or a chat "
        "message. The launch readiness score is calculated server-side by the "
        "lc_CalculateLaunchReadiness Custom API in Dataverse, which has full "
        "visibility of every milestone and task. The agent must invoke that API "
        "and report its three return values verbatim: lc_ReadinessScore "
        "(decimal 0–100), lc_Verdict (GO | CONDITIONAL | NO-GO), and "
        "lc_ReadinessSummary (a per-milestone narrative)."
    )
    doc.add_paragraph(
        "Verdict thresholds: score ≥ 85 is GO. Score between 65 and 84 inclusive "
        "is CONDITIONAL — leadership sign-off required. Below 65 is NO-GO."
    )

    doc.add_heading("1.2 Resolving the launch in question", level=2)
    doc.add_paragraph(
        "If the user names a launch (\"Q3 Widget Launch\"), use that. If they "
        "don't name one, query lc_launch and ask which one — do not guess and "
        "do not pick the most recent one silently."
    )

    doc.add_heading("1.3 Reporting the result", level=2)
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("State the verdict in the first sentence.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Quote the score with one decimal place.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Surface any milestone that is At Risk or Blocked, by name.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("If verdict is CONDITIONAL, name who needs to sign off.")

    doc.add_heading("2. Escalation when things go sideways", level=1)
    doc.add_paragraph(
        "A task or milestone marked Blocked is an escalation candidate. Not every "
        "blocker requires a page; we have killed enough Saturdays to know the "
        "difference. The agent's job is to decide severity, then route."
    )

    doc.add_heading("2.1 Verify the block before paging anyone (POLICY)", level=2)
    doc.add_paragraph(
        "Before any escalation, the agent cross-checks Dataverse against the "
        "linked GitHub issue (we model GitHub Issues as a virtual entity — "
        "lc_GitHubIssueId on lc_task). If a task is marked Blocked in Dataverse "
        "but the linked GitHub issue is closed, that is stale data, not a real "
        "blocker. DO NOT escalate. Recommend transitioning the task back to In "
        "Progress and re-running readiness."
    )

    doc.add_heading("2.2 Severity scoring", level=2)
    doc.add_paragraph(
        "Severity is one of Low, Medium, High, or Critical. Inputs:"
    )
    doc.add_paragraph(
        "GitHub label signal — if the linked GitHub issue carries a \"blocker\" "
        "or \"P0\" label, severity is at minimum High, possibly Critical. "
        "GitHub labels are a stronger signal than Dataverse blocker text alone "
        "because engineering tends to triage there first.", style="List Bullet"
    )
    doc.add_paragraph(
        "Milestone proximity — a blocker within 7 days of the target date is "
        "escalated one level. Within 2 days is escalated two levels. Critical "
        "is the ceiling.", style="List Bullet"
    )
    doc.add_paragraph(
        "Cross-team dependency — if the blocker requires another team to act, "
        "severity floor is Medium even if the proximity is far. Cross-team "
        "blockers rot if they are not visible.", style="List Bullet"
    )

    doc.add_heading("2.3 Notification chain", level=2)
    doc.add_paragraph("Low — log only, no page. Mention in the next standup.")
    doc.add_paragraph(
        "Medium — message the task owner directly. If no acknowledgement within "
        "4 business hours, escalate to the milestone owner."
    )
    doc.add_paragraph(
        "High — page the task owner AND the milestone owner. Same channel. "
        "Expectation is acknowledgement within 1 hour."
    )
    doc.add_paragraph(
        "Critical — page the launch owner, the milestone owner, and the on-call. "
        "Open a war room channel. Acknowledgement expected within 15 minutes. "
        "If the launch is in CONDITIONAL or NO-GO state and a Critical blocker "
        "lands, the launch owner must trigger a go/no-go re-review within 24 "
        "hours regardless of where we are in the cycle."
    )

    doc.add_heading("3. Status transitions — which arrows are legal", level=1)
    doc.add_paragraph(
        "Status changes look innocuous in a UI but they are policy decisions. "
        "Each entity has its own status column with its own legal transitions. "
        "The agent must respect these. Do NOT free-text statuses, and do NOT "
        "operate on the staging tracker tables (lc_trackera..lc_trackere) for "
        "live status questions — those are append-only landing zones."
    )

    doc.add_heading("3.1 Canonical status columns (POLICY)", level=2)
    doc.add_paragraph(
        "Launch status lives on lc_launch.lc_launchstatus. Milestone status "
        "lives on lc_milestone.lc_milestonestatus. Task status lives on "
        "lc_task.lc_taskstatus. The option-set integer codes are entity-specific. "
        "Mixing them up will yield HTTP 400 and a lot of confusion."
    )

    doc.add_heading("3.2 Launch transitions", level=2)
    doc.add_paragraph(
        "Planning → In Progress → Ready for Launch → Launched is the happy path. "
        "From any state except Launched and Cancelled, the launch may move to On "
        "Hold. From On Hold the launch resumes to whichever state it was in. A "
        "launch may move to Cancelled from any non-terminal state, but cancellation "
        "requires a written reason in the status notes."
    )

    doc.add_heading("3.3 Milestone transitions", level=2)
    doc.add_paragraph(
        "Not Started → In Progress → Complete is the happy path. A milestone may "
        "move into At Risk from In Progress when any of its tasks slip past their "
        "due date. A milestone moves to Blocked if any of its tasks are Blocked. "
        "Once Complete, a milestone is terminal — do not reopen. If you discover "
        "additional work, file a new milestone."
    )

    doc.add_heading("3.4 Task transitions", level=2)
    doc.add_paragraph(
        "Not Started → In Progress → Done. From In Progress, a task may move to "
        "Blocked at any time; from Blocked it returns to In Progress when the "
        "blocker clears. Done is terminal."
    )

    doc.add_heading("4. Notes from previous reviews", level=1)
    doc.add_paragraph(
        "[Q1] Marketing asked whether the readiness verdict could be overridden "
        "if a single milestone is At Risk but recoverable. Answer: no, but the "
        "CONDITIONAL verdict exists for exactly this case. Treat CONDITIONAL as "
        "GO with leadership sign-off and a documented mitigation plan."
    )
    doc.add_paragraph(
        "[Q2] Engineering asked us to stop trusting the Dataverse blocked flag "
        "without checking GitHub. We agreed. See section 2.1."
    )
    doc.add_paragraph(
        "[Q3] Add the Custom API note (1.1) after the Q3 retro found three "
        "launches where someone tallied gates by hand and got it wrong."
    )

    p = doc.add_paragraph()
    p.add_run("— end of playbook —").italic = True

    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print(f"wrote {path}")
